import random
from enum import Enum
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt,RGBColor
from docx.oxml.ns import qn
from docx import Document


class Expression(object):
    def __init__(self):
        pass

    def generate(self):
        pass


class Mul(Expression):
    def __init__(self):
        super().__init__()
        self.__r = 1
        self.__l = 1

    @staticmethod
    def name():
        return "乘"

    def generate(self):
        self.__l = random.randint(2, 9)
        self.__r = random.randint(2, 9)
        pass

    def __str__(self):
        return f'{self.__l} × {self.__r} = '

    pass


class Div(Expression):
    def __init__(self):
        super().__init__()
        self.__r = 1
        self.__l = 1

    @staticmethod
    def name():
        return "除"

    def generate(self):
        a = random.randint(2, 9)
        b = random.randint(2, 9)
        c = a * b
        self.__l = c
        self.__r = a
        pass

    def __str__(self):
        return f'{self.__l} ÷ {self.__r} = '

    pass


class Add(Expression):
    def __init__(self):
        super().__init__()
        self.__r = 1
        self.__l = 1

    @staticmethod
    def name():
        return "加"

    def generate(self):
        self.__l = random.randint(1, 99)
        self.__r = random.randint(1, 100-self.__l)
        pass

    def __str__(self):
        return f'{self.__l} + {self.__r} = '

    pass


class Sub(Expression):
    def __init__(self):
        super().__init__()
        self.__r = 1
        self.__l = 1

    @staticmethod
    def name():
        return "减"

    def generate(self):
        a = random.randint(1, 99)
        b = random.randint(1, 99)
        self.__l = max(a, b)
        self.__r = min(a, b)
        pass

    def __str__(self):
        return f'{self.__l} - {self.__r} = '


class ExpressionFactory(object):
    def __init__(self):
        self.opCreatorSet = {}

    def register(self, op, creator):
        self.opCreatorSet[op] = creator

    def get_creator(self, op):
        return self.opCreatorSet[op]


class Sheet(object):
    def __init__(self, operators: list = []):

        ef = ExpressionFactory()
        ef.register("+", Add)
        ef.register("-", Sub)
        ef.register("*", Mul)
        ef.register("/", Div)

        self.__expArr = []
        self.__creators = []
        for op in operators:
            self.__creators.append(ef.get_creator(op))

    def generate_exp(self, n) -> list:
        res = []
        cache = {}
        for i in range(n):
            idx = random.randint(0, len(self.__creators)-1)
            e = self.__creators[idx]()
            e.generate()
            has = cache.get(str(e), False)
            while has :
                e.generate()
                has = cache.get(str(e), False)
            cache[str(e)] = True
            res.append(str(e))
        return res

    def get_title(self,index):
        name = ""
        for op in self.__creators:
            name += op().name()
        return f'100以内整数{name}练习({index})'

    def generate(self,  page_cnt):

        document = Document()
        document.styles['Normal'].font.name = u'Arial'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Arial')
        document.styles['Normal'].font.size = Pt(16)
        document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

        for i in range(page_cnt):
            exps = self.generate_exp(57)
            for e in exps:
                print(e)

            if i != 0:
                sec = document.add_section(start_type=WD_SECTION.CONTINUOUS)
                sec._sectPr.xpath('./w:cols')[0].set(qn('w:num'),'1')

            p = document.add_heading(self.get_title(i+1), 0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            sec = document.add_section(start_type=WD_SECTION.CONTINUOUS)
            sec._sectPr.xpath('./w:cols')[0].set(qn('w:num'),'3')
            for e in exps:
                p = document.add_paragraph(e)

        document.save("/Users/cuihandong/Desktop/数学加减法练习.docx")


if __name__ == '__main__':

    # sh = Sheet(["+", "*", "/", "-"])
    sh = Sheet(["+", "-"])
    sh.generate(10)
